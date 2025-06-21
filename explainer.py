import os
import re
import xml.etree.ElementTree as ET
import glob
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
from PIL import Image, ImageGrab
import logging
from datetime import datetime
import sys
from pathlib import Path

# Try to import optional PDF conversion library
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
    print("PDF conversion available via docx2pdf")
except ImportError:
    PDF_AVAILABLE = False
    print("PDF conversion not available. Install docx2pdf if you need PDF output: pip install docx2pdf")

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='alteryx_workflow_documentation.log'
)

class AlteryxWorkflowDocumenter:
    """Class to document Alteryx workflows for audit purposes"""
    
    def __init__(self, workflow_dir="C:\\Users\\framework"):
        """Initialize the documenter with the directory containing workflow files"""
        self.workflow_dir = workflow_dir
        self.doc = Document()
        self.workflows = []
        
        # Enhanced component descriptions for auditors
        self.component_descriptions = {
            "Input": {
                "purpose": "Reads data from external sources",
                "audit_concern": "Data source integrity, access controls, data lineage",
                "what_it_does": "This component connects to and retrieves data from files, databases, or other data sources. It's the starting point of the data processing workflow."
            },
            "Output": {
                "purpose": "Writes processed data to destination",
                "audit_concern": "Output location security, data retention, access permissions",
                "what_it_does": "This component saves the final processed data to a specified location such as a file, database, or report. This is where the workflow results are stored."
            },
            "Filter": {
                "purpose": "Removes unwanted records based on criteria",
                "audit_concern": "Business logic accuracy, completeness of filtering rules",
                "what_it_does": "This component examines each data record and keeps only those that meet specific conditions. Records that don't meet the criteria are excluded from further processing."
            },
            "Formula": {
                "purpose": "Creates new fields or modifies existing data using calculations",
                "audit_concern": "Calculation accuracy, business rule implementation",
                "what_it_does": "This component performs mathematical calculations, text manipulations, or logical operations to create new data fields or modify existing ones."
            },
            "Join": {
                "purpose": "Combines data from multiple sources based on common fields",
                "audit_concern": "Data matching accuracy, handling of unmatched records",
                "what_it_does": "This component merges data from two or more sources by matching records that have the same values in specified fields (like matching customer IDs)."
            },
            "Summarize": {
                "purpose": "Aggregates data by grouping and calculating totals, averages, counts",
                "audit_concern": "Aggregation logic accuracy, completeness of groupings",
                "what_it_does": "This component groups related records together and calculates summary statistics like totals, counts, averages, or other aggregate functions."
            },
            "Sort": {
                "purpose": "Arranges data in a specific order",
                "audit_concern": "Sorting consistency, impact on downstream processing",
                "what_it_does": "This component reorders the data records based on specified fields, either in ascending or descending order."
            },
            "Union": {
                "purpose": "Combines datasets by stacking records from multiple sources",
                "audit_concern": "Data structure consistency, duplicate handling",
                "what_it_does": "This component stacks data from multiple sources on top of each other, creating one larger dataset with all the records."
            },
            "Select": {
                "purpose": "Chooses specific fields and controls data structure",
                "audit_concern": "Data completeness, field selection rationale",
                "what_it_does": "This component selects which data fields to keep, rename, or reorder. It's used to clean up the data structure and keep only relevant information."
            },
            "Browse": {
                "purpose": "Displays data for review and quality checking",
                "audit_concern": "Data quality validation points",
                "what_it_does": "This component allows users to view the data at specific points in the workflow for quality checking and validation purposes."
            }
        }
        logging.info(f"Initialized AlteryxWorkflowDocumenter with directory: {workflow_dir}")
    
    def find_workflow_files(self):
        """Find all Alteryx workflow files in the specified directory"""
        print(f"Looking for workflow files in: {self.workflow_dir}")
        
        # Check if directory exists
        if not os.path.exists(self.workflow_dir):
            print(f"ERROR: Directory {self.workflow_dir} does not exist!")
            logging.error(f"Directory {self.workflow_dir} does not exist")
            return []
        
        # List all files in the directory for debugging
        try:
            all_files = os.listdir(self.workflow_dir)
            print(f"All files in directory: {all_files}")
        except Exception as e:
            print(f"Error listing directory contents: {e}")
            return []
        
        # Look for .yxmd files
        workflow_pattern = os.path.join(self.workflow_dir, "*.yxmd")
        yxmd_files = glob.glob(workflow_pattern)
        print(f"Found .yxmd files: {yxmd_files}")
        
        # Also look for .bak files (backup Alteryx files)
        backup_pattern = os.path.join(self.workflow_dir, "*.bak")
        bak_files = glob.glob(backup_pattern)
        print(f"Found .bak files: {bak_files}")
        
        # Combine both types
        self.workflows = yxmd_files + bak_files
        
        print(f"Total workflow files found: {len(self.workflows)}")
        logging.info(f"Found {len(self.workflows)} workflow files")
        
        return self.workflows
    
    def parse_workflow(self, workflow_path):
        """Parse an Alteryx workflow file and extract component information"""
        try:
            print(f"Parsing workflow: {workflow_path}")
            tree = ET.parse(workflow_path)
            root = tree.getroot()
            
            # Extract workflow metadata
            workflow_name = os.path.basename(workflow_path)
            workflow_info = {
                'name': workflow_name,
                'path': workflow_path,
                'components': [],
                'connections': []
            }
            
            # Extract components (nodes)
            nodes = root.findall(".//Node")
            print(f"Found {len(nodes)} nodes in workflow")
            
            for node in nodes:
                component = {
                    'type': node.get('ToolID', ''),
                    'plugin': node.get('Plugin', ''),
                    'id': node.get('ToolID', ''),
                    'label': '',
                    'description': '',
                    'inputs': [],
                    'outputs': [],
                    'properties': {},
                    'position': {'x': 0, 'y': 0}
                }
                
                # Extract position for workflow visualization
                gui_settings = node.find(".//GuiSettings")
                if gui_settings is not None:
                    component['label'] = gui_settings.get('Html', '')
                    position = gui_settings.get('Position', '0,0').split(',')
                    if len(position) >= 2:
                        try:
                            component['position']['x'] = int(position[0])
                            component['position']['y'] = int(position[1])
                        except:
                            pass
                
                # Extract properties with better organization
                properties = node.findall(".//Properties/Configuration/*")
                for prop in properties:
                    if prop.text:
                        component['properties'][prop.tag] = prop.text
                
                # Extract connections
                connections = root.findall(f".//Connection[@TargetID='{component['id']}']")
                for conn in connections:
                    component['inputs'].append({
                        'source_id': conn.get('SourceID', ''),
                        'source_type': '',
                        'name': conn.get('Name', ''),
                        'wireless': conn.get('Wireless', 'False') == 'True'
                    })
                
                connections = root.findall(f".//Connection[@SourceID='{component['id']}']")
                for conn in connections:
                    component['outputs'].append({
                        'target_id': conn.get('TargetID', ''),
                        'target_type': '',
                        'name': conn.get('Name', ''),
                        'wireless': conn.get('Wireless', 'False') == 'True'
                    })
                
                workflow_info['components'].append(component)
            
            # Fill in source/target types and create flow map
            component_dict = {comp['id']: comp for comp in workflow_info['components']}
            for component in workflow_info['components']:
                for input_conn in component['inputs']:
                    if input_conn['source_id'] in component_dict:
                        source_comp = component_dict[input_conn['source_id']]
                        source_type = source_comp['plugin'].split('.')[-1] if '.' in source_comp['plugin'] else source_comp['plugin']
                        input_conn['source_type'] = source_type
                
                for output_conn in component['outputs']:
                    if output_conn['target_id'] in component_dict:
                        target_comp = component_dict[output_conn['target_id']]
                        target_type = target_comp['plugin'].split('.')[-1] if '.' in target_comp['plugin'] else target_comp['plugin']
                        output_conn['target_type'] = target_type
            
            logging.info(f"Successfully parsed workflow: {workflow_name}")
            return workflow_info
        except Exception as e:
            print(f"Error parsing workflow {workflow_path}: {str(e)}")
            logging.error(f"Error parsing workflow {workflow_path}: {str(e)}")
            return None
    
    def analyze_workflow_flow(self, workflow_info):
        """Analyze the workflow to understand the data flow sequence"""
        components = workflow_info['components']
        component_dict = {comp['id']: comp for comp in components}
        
        # Find input components (no inputs)
        inputs = [comp for comp in components if not comp['inputs']]
        
        # Find output components (no outputs)
        outputs = [comp for comp in components if not comp['outputs']]
        
        # Build flow sequence
        flow_sequence = []
        visited = set()
        
        def trace_flow(component, path=[]):
            if component['id'] in visited:
                return
            visited.add(component['id'])
            
            current_path = path + [component]
            
            if not component['outputs']:  # This is an end point
                flow_sequence.append(current_path)
            else:
                for output in component['outputs']:
                    if output['target_id'] in component_dict:
                        next_comp = component_dict[output['target_id']]
                        trace_flow(next_comp, current_path)
        
        # Start tracing from input components
        for input_comp in inputs:
            trace_flow(input_comp)
        
        return {
            'inputs': inputs,
            'outputs': outputs,
            'flow_sequences': flow_sequence
        }
    
    def generate_audit_documentation(self, workflow_info):
        """Generate comprehensive audit documentation"""
        try:
            # Create a new document with better formatting
            self.doc = Document()
            
            # Set up styles
            styles = self.doc.styles
            
            # Add title page
            title = self.doc.add_heading('ALTERYX WORKFLOW AUDIT DOCUMENTATION', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.doc.add_paragraph('')
            
            # Add workflow identification
            self.doc.add_paragraph(f'Workflow Name: {workflow_info["name"]}', style='Heading 2')
            self.doc.add_paragraph(f'File Location: {workflow_info["path"]}')
            self.doc.add_paragraph(f'Documentation Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
            self.doc.add_paragraph(f'Total Components: {len(workflow_info["components"])}')
            
            # Page break
            self.doc.add_page_break()
            
            # Executive Summary
            self.doc.add_heading('EXECUTIVE SUMMARY', level=1)
            
            # Analyze workflow
            flow_analysis = self.analyze_workflow_flow(workflow_info)
            
            summary_text = f"""
This document provides a comprehensive audit trail for the Alteryx workflow '{workflow_info["name"]}'. 
This workflow contains {len(workflow_info["components"])} components that process data through {len(flow_analysis['flow_sequences'])} main processing paths.

KEY COMPONENTS:
• {len(flow_analysis['inputs'])} Data Input Sources
• {len(flow_analysis['outputs'])} Output Destinations  
• {len([c for c in workflow_info["components"] if 'Filter' in c.get('plugin', '')])} Data Filters
• {len([c for c in workflow_info["components"] if 'Formula' in c.get('plugin', '')])} Calculation/Formula Components
• {len([c for c in workflow_info["components"] if 'Join' in c.get('plugin', '')])} Data Joins
• {len([c for c in workflow_info["components"] if 'Summarize' in c.get('plugin', '')])} Aggregation/Summary Components

This workflow is designed to process data through a series of validated steps, ensuring data quality and business rule compliance at each stage.
            """
            self.doc.add_paragraph(summary_text.strip())
            
            # Data Flow Overview
            self.doc.add_heading('DATA FLOW OVERVIEW', level=1)
            
            self.doc.add_paragraph("This section describes how data moves through the workflow from start to finish:")
            
            for i, flow_path in enumerate(flow_analysis['flow_sequences'], 1):
                self.doc.add_heading(f'Processing Path {i}', level=2)
                
                for j, component in enumerate(flow_path, 1):
                    component_type = component['plugin'].split('.')[-1] if '.' in component['plugin'] else component['plugin']
                    
                    self.doc.add_paragraph(f"Step {j}: {component_type}", style='Heading 3')
                    
                    # Get component description
                    comp_info = self.component_descriptions.get(component_type, {
                        "purpose": f"Performs {component_type} operations",
                        "what_it_does": f"This component handles {component_type} functionality in the workflow."
                    })
                    
                    self.doc.add_paragraph(f"Purpose: {comp_info['purpose']}")
                    self.doc.add_paragraph(f"What it does: {comp_info['what_it_does']}")
                    
                    # Add key properties if available
                    if component['properties']:
                        key_props = {k: v for k, v in component['properties'].items() 
                                   if k in ['FileName', 'File', 'OutputMode', 'Expression', 'Filter']}
                        if key_props:
                            self.doc.add_paragraph("Key Configuration:")
                            for prop, value in key_props.items():
                                if len(str(value)) > 100:
                                    value = str(value)[:100] + "..."
                                self.doc.add_paragraph(f"• {prop}: {value}")
                    
                    self.doc.add_paragraph("")
            
            # Detailed Component Analysis
            self.doc.add_heading('DETAILED COMPONENT ANALYSIS', level=1)
            
            # Group components by type for better organization
            components_by_type = {}
            for component in workflow_info['components']:
                component_type = component['plugin'].split('.')[-1] if '.' in component['plugin'] else component['plugin']
                if component_type not in components_by_type:
                    components_by_type[component_type] = []
                components_by_type[component_type].append(component)
            
            for component_type, components in components_by_type.items():
                self.doc.add_heading(f'{component_type} Components ({len(components)})', level=2)
                
                # Add component type description
                comp_info = self.component_descriptions.get(component_type, {})
                if comp_info:
                    self.doc.add_paragraph(f"Purpose: {comp_info.get('purpose', 'N/A')}")
                    if 'audit_concern' in comp_info:
                        self.doc.add_paragraph(f"Audit Focus: {comp_info['audit_concern']}")
                    self.doc.add_paragraph("")
                
                # Document each component of this type
                for i, component in enumerate(components, 1):
                    self.doc.add_heading(f'{component_type} #{i}: {component.get("label", component["id"])}', level=3)
                    
                    # Component details
                    details_table = self.doc.add_table(rows=1, cols=2)
                    details_table.style = 'Table Grid'
                    header_cells = details_table.rows[0].cells
                    header_cells[0].text = 'Attribute'
                    header_cells[1].text = 'Value'
                    
                    # Add component ID
                    row_cells = details_table.add_row().cells
                    row_cells[0].text = 'Component ID'
                    row_cells[1].text = component['id']
                    
                    # Add inputs if any
                    if component['inputs']:
                        row_cells = details_table.add_row().cells
                        row_cells[0].text = 'Data Sources'
                        sources = ', '.join([f"{inp['source_type']} ({inp['source_id']})" for inp in component['inputs']])
                        row_cells[1].text = sources
                    
                    # Add outputs if any
                    if component['outputs']:
                        row_cells = details_table.add_row().cells
                        row_cells[0].text = 'Data Destinations'
                        destinations = ', '.join([f"{out['target_type']} ({out['target_id']})" for out in component['outputs']])
                        row_cells[1].text = destinations
                    
                    # Add key properties
                    important_props = ['FileName', 'File', 'OutputMode', 'Expression', 'Filter', 'Mode']
                    for prop in important_props:
                        if prop in component['properties'] and component['properties'][prop]:
                            row_cells = details_table.add_row().cells
                            row_cells[0].text = prop
                            prop_value = str(component['properties'][prop])
                            if len(prop_value) > 200:
                                prop_value = prop_value[:200] + "... (truncated)"
                            row_cells[1].text = prop_value
                    
                    self.doc.add_paragraph("")
            
            # Quality Control and Validation Points
            self.doc.add_heading('QUALITY CONTROL & VALIDATION POINTS', level=1)
            
            validation_text = """
The following components serve as quality control and validation points in the workflow:

BROWSE COMPONENTS: These components allow for data inspection at key points in the workflow. 
During execution, users can review data quality, record counts, and field values to ensure 
processing is occurring correctly.

FILTER COMPONENTS: These components implement business rules and data quality checks by 
removing invalid or unwanted records. The filter conditions should be reviewed to ensure 
they align with business requirements.

OUTPUT COMPONENTS: These components represent the final deliverables of the workflow. 
The output locations and formats should be verified for security and access control compliance.
            """
            self.doc.add_paragraph(validation_text.strip())
            
            # Browse components
            browse_components = [c for c in workflow_info['components'] if 'Browse' in c.get('plugin', '')]
            if browse_components:
                self.doc.add_heading('Data Review Points', level=2)
                for browse in browse_components:
                    self.doc.add_paragraph(f"• Review Point {browse['id']}: Allows inspection of data after processing steps")
            
            # Filter components
            filter_components = [c for c in workflow_info['components'] if 'Filter' in c.get('plugin', '')]
            if filter_components:
                self.doc.add_heading('Data Quality Filters', level=2)
                for filt in filter_components:
                    filter_expr = filt['properties'].get('Filter', 'Not specified')
                    self.doc.add_paragraph(f"• Filter {filt['id']}: {filter_expr}")
            
            # Testing and Replication Instructions
            self.doc.add_heading('TESTING AND REPLICATION INSTRUCTIONS', level=1)
            
            testing_instructions = f"""
To test or replicate this workflow:

1. ENVIRONMENT SETUP:
   • Ensure Alteryx Designer is installed and licensed
   • Verify access to all input data sources identified in the Input components
   • Confirm write permissions to all output destinations

2. DATA PREPARATION:
   • Prepare test data that matches the structure of the original input sources
   • Ensure test data includes edge cases and boundary conditions
   • Document any data dependencies or prerequisites

3. EXECUTION STEPS:
   • Open the workflow file: {workflow_info['name']}
   • Review all Input components to ensure data source connections are valid
   • Run the workflow in segments using Browse components to validate intermediate results
   • Compare output results with expected outcomes
   • Document any discrepancies or errors

4. VALIDATION CHECKLIST:
   • Verify record counts at each major processing step
   • Check calculation accuracy in Formula components  
   • Validate filter logic removes appropriate records
   • Confirm output format and location match requirements
   • Test with various data scenarios (normal, edge cases, empty datasets)
            """
            self.doc.add_paragraph(testing_instructions.strip())
            
            # Access Control and Security
            self.doc.add_heading('ACCESS CONTROL AND SECURITY', level=1)
            
            security_text = """
WORKFLOW FILE ACCESS:
The workflow file itself should be stored in a secure location with appropriate access controls. 
Only authorized personnel should have the ability to modify the workflow logic.

DATA SOURCE SECURITY:
Review the security settings and access controls for all input data sources. Ensure that 
data access follows the principle of least privilege.

OUTPUT DESTINATION SECURITY:  
Verify that output files and databases have appropriate security settings. Consider data 
sensitivity and implement encryption where required.

AUDIT TRAIL:
Maintain logs of workflow execution, including who ran the workflow, when it was executed, 
and any errors or warnings that occurred during processing.
            """
            self.doc.add_paragraph(security_text.strip())
            
            # Input and Output Summary
            self.doc.add_heading('INPUT AND OUTPUT SUMMARY', level=1)
            
            # Input sources
            input_components = [c for c in workflow_info['components'] if 'Input' in c.get('plugin', '')]
            if input_components:
                self.doc.add_heading('Data Input Sources', level=2)
                for inp in input_components:
                    file_path = inp['properties'].get('FileName', inp['properties'].get('File', 'Not specified'))
                    self.doc.add_paragraph(f"• Input {inp['id']}: {file_path}")
            
            # Output destinations  
            output_components = [c for c in workflow_info['components'] if 'Output' in c.get('plugin', '')]
            if output_components:
                self.doc.add_heading('Data Output Destinations', level=2)
                for out in output_components:
                    file_path = out['properties'].get('FileName', out['properties'].get('File', 'Not specified'))
                    self.doc.add_paragraph(f"• Output {out['id']}: {file_path}")
            
            logging.info(f"Successfully generated audit documentation for workflow: {workflow_info['name']}")
            return True
            
        except Exception as e:
            logging.error(f"Error generating documentation: {str(e)}")
            print(f"Error generating documentation: {str(e)}")
            return False
    
    def save_documentation(self, output_filename):
        """Save the documentation as Word document and optionally as PDF"""
        try:
            # Save as Word document
            word_filename = f"{output_filename}.docx"
            self.doc.save(word_filename)
            logging.info(f"Saved Word document: {word_filename}")
            print(f"Word document saved: {word_filename}")
            
            # Try to convert to PDF if library is available
            if PDF_AVAILABLE:
                try:
                    pdf_filename = f"{output_filename}.pdf"
                    convert(word_filename, pdf_filename)
                    logging.info(f"Saved PDF document: {pdf_filename}")
                    print(f"PDF document saved: {pdf_filename}")
                except Exception as pdf_error:
                    logging.error(f"Error converting to PDF: {str(pdf_error)}")
                    print(f"Warning: Could not create PDF. Error: {pdf_error}")
            else:
                print("PDF conversion skipped (docx2pdf not installed)")
            
            return word_filename
        except Exception as e:
            logging.error(f"Error saving documentation: {str(e)}")
            print(f"Error saving documentation: {str(e)}")
            return None

def main():
    """Main function to run the workflow documentation process"""
    print("Alteryx Workflow Audit Documentation Generator")
    print("==============================================")
    
    # Get the current working directory and check for workflows there too
    current_dir = os.getcwd()
    print(f"Current working directory: {current_dir}")
    
    # Let user specify directory or use default
    workflow_dir = input("Enter the directory path containing workflow files (or press Enter for C:\\Users\\framework): ").strip()
    if not workflow_dir:
        workflow_dir = "C:\\Users\\framework"
    
    # Initialize the documenter
    documenter = AlteryxWorkflowDocumenter(workflow_dir)
    
    # Find workflow files
    workflows = documenter.find_workflow_files()
    if not workflows:
        print("No workflow files found in the specified directory.")
        print("Make sure you have .yxmd or .bak files in the directory.")
        return
    
    print(f"Found {len(workflows)} workflow files.")
    
    for i, workflow_path in enumerate(workflows):
        print(f"\nProcessing workflow {i+1}/{len(workflows)}: {os.path.basename(workflow_path)}")
        
        # Parse the workflow
        workflow_info = documenter.parse_workflow(workflow_path)
        if workflow_info is None:
            print(f"Error parsing workflow. Check the log for details.")
            continue
        
        # Generate audit documentation
        if documenter.generate_audit_documentation(workflow_info):
            # Save documentation
            output_base = os.path.splitext(os.path.basename(workflow_path))[0]
            output_filename = f"{output_base}_audit_documentation"
            saved_file = documenter.save_documentation(output_filename)
            
            if saved_file:
                print(f"Audit documentation saved to {saved_file}")
                if PDF_AVAILABLE:
                    print(f"PDF version also created.")
                else:
                    print(f"Only Word version created (install docx2pdf for PDF support).")
            else:
                print("Error saving documentation. Check the log for details.")
        else:
            print("Error generating documentation. Check the log for details.")
    
    print("\nAudit documentation process complete.")

if __name__ == "__main__":
    main()
